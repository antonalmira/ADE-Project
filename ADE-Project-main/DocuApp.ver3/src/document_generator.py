import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.shared import OxmlElement, qn
from PyQt5.QtCore import Qt
from performance_section import PerformanceSection
from waveform_section import WaveformSection
from utils import ensure_directory, remove_directory

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_inner_borders(table, hex_color):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:color'), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
def set_table_all_borders(table, hex_color):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:color'), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def format_text_specs(text):
    if not isinstance(text, str):
        return text
    text = re.sub(r'(?<=\d)(?=[a-zA-ZµΩ°])', ' ', text)
    for term in ['vac', 'vdc', 'vor', 'kp']:
        text = re.sub(fr'\b{term}\b', term.upper(), text, flags=re.IGNORECASE)
    return text

def apply_column_widths(table, width_inches_list):
    for row in table.rows:
        for idx, width in enumerate(width_inches_list):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)

def add_bom_table(document, excel_path):
    try:
        xl = pd.ExcelFile(excel_path)
        # Check specifically for Mechanical and Electrical tables 
        sheets_to_process = [s for s in xl.sheet_names if 'Electrical' in s or 'Mechanical' in s]
        
        if not sheets_to_process:
            print("No 'Electrical' or 'Mechanical' sheets found in BOM.")
            return

        for sheet_name in sheets_to_process:
            df = pd.read_excel(excel_path, sheet_name=sheet_name, skiprows=2)
            bom_columns = ['Item', 'Quantity', 'Designator', 'Value', 'Description', 'Manufacturer Part Number', 'Manufacturer']
            
            rename_map = {}
            for col in df.columns:
                if 'part number' in str(col).lower() and 'man' in str(col).lower(): rename_map[col] = 'Manufacturer Part Number'
                elif 'man' in str(col).lower() and 'part' not in str(col).lower(): rename_map[col] = 'Manufacturer'
            df = df.rename(columns=rename_map)
            
            existing_cols = [col for col in bom_columns if col in df.columns]
            df = df[existing_cols].dropna(how='all')
            
            target_keywords = ["electrical"] if "Electrical" in sheet_name else ["mechanical"]
            target_para = None
            
            # Step 1: Search for specific 14.1 or 14.2 heading
            for p in document.paragraphs:
                if any(k in p.text.lower() for k in target_keywords) and ("14." in p.text or "bill of materials" in p.text.lower()):
                    target_para = p
                    break
                    
            # Step 2: Fallback to main BOM heading if specific one isn't found
            if not target_para:
                for p in document.paragraphs:
                    if "bill of materials" in p.text.lower():
                        target_para = p
                        break
                        
            # Step 3: If no section exists at all, just create a header at the end of the doc
            if not target_para:
                target_para = document.add_paragraph(f"{sheet_name} Bill of Materials", style='Heading 2')

            table = document.add_table(rows=1, cols=len(df.columns))
            set_table_inner_borders(table, 'C0C0C0')
            
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            table.rows[0].height = Pt(35)

            for i, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[i]
                cell.text = str(col_name)
                set_cell_background(cell, '0085CA') 
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    cell_val = "" if pd.isna(value) else str(value)
                    col_name = df.columns[i]
                    
                    if col_name == 'Designator':
                        cell_val = cell_val.replace(',', '\n').replace(' ', '')
                    else:
                        cell_val = format_text_specs(cell_val)

                    cell = row_cells[i]
                    cell.text = cell_val
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    p = cell.paragraphs[0]
                    if col_name == 'Description':
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)

            apply_column_widths(table, [0.4, 0.6, 0.72, 0.6, 2.16, 1.65, 1.01])
            
            # Reposition the table to sit below the actual section
            tbl_element = table._element
            tbl_element.getparent().remove(tbl_element)
            target_para._element.addnext(tbl_element)

    except Exception as e:
        print(f"BOM Error: {e}")

def add_pixls_designer_table(document, excel_path):
    try:
        xl = pd.ExcelFile(excel_path)
        sheet_name = next((s for s in xl.sheet_names if 'Design' in s), None)
        if not sheet_name: return

        df_raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
        header_idx = 0
        for i, row in df_raw.iterrows():
            row_str = [str(x).upper() for x in row.values]
            if 'INPUT' in row_str and 'OUTPUT' in row_str:
                header_idx = i
                break

        df = pd.read_excel(excel_path, sheet_name=sheet_name, skiprows=header_idx).iloc[:, :7]
        raw_headers = df.columns.tolist()
        clean_headers = ["" if "Unnamed" in str(col) else str(col) for col in raw_headers]
        df.columns = ['Row Num', 'Parameter Name', 'INPUT', 'INFO', 'OUTPUT', 'UNIT', 'Description']
        
        target_para = None
        for p in document.paragraphs:
            if "design spreadsheet" in p.text.lower():
                target_para = p
                break
                
        if not target_para:
            target_para = document.add_paragraph("Design Spreadsheet", style='Heading 2')

        table = document.add_table(rows=1, cols=len(df.columns))
        set_table_all_borders(table, 'C0C0C0') 

        for i, col_name in enumerate(clean_headers):
            cell = table.rows[0].cells[i]
            cell.text = format_text_specs(col_name)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, 'FFFFFF') 
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0) 

        for _, row in df.iterrows():
            if pd.isna(row['Parameter Name']) and pd.isna(row['INPUT']): continue
            row_cells = table.add_row().cells
            is_subheader = pd.isna(row['INPUT']) and pd.isna(row['OUTPUT']) and pd.isna(row['UNIT'])

            if is_subheader:
                cell_0 = row_cells[0]
                num_val = str(row['Row Num']) if not pd.isna(row['Row Num']) else ""
                if num_val.endswith('.0'): num_val = num_val[:-2] 
                cell_0.text = num_val
                cell_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p0 = cell_0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r0 = p0.runs[0]
                r0.font.name = 'Calibri'
                r0.font.size = Pt(8)

                main_cell = row_cells[1]
                main_cell.merge(row_cells[5]) 
                main_cell.text = format_text_specs(str(row['Parameter Name']))
                set_cell_background(main_cell, 'D9D9D9') 
                main_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                p = main_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255) 

                desc_cell = row_cells[6]
                desc_val = "" if pd.isna(row['Description']) else format_text_specs(str(row['Description']))
                desc_cell.text = desc_val
                set_cell_background(desc_cell, 'FFFFFF') 
                desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                p_desc = desc_cell.paragraphs[0]
                p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p_desc.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)

            else:
                for i, col_name in enumerate(df.columns):
                    val = "" if pd.isna(row[col_name]) else format_text_specs(str(row[col_name]))
                    if col_name == 'Row Num' and val.endswith('.0'): 
                        val = val[:-2] 

                    cell = row_cells[i]
                    cell.text = val
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    p = cell.paragraphs[0]
                    if col_name in ['Parameter Name', 'Description']:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)

        apply_column_widths(table, [0.32, 1.37, 0.7, 0.7, 0.7, 0.5, 2.52])
        
        # Reposition the table below the targeted section header
        tbl_element = table._element
        tbl_element.getparent().remove(tbl_element)
        target_para._element.addnext(tbl_element)

    except Exception as e:
        print(f"PIXls Error: {e}")

class DocGenerator:
    def __init__(self, app, output_path, update_doc_path=""):
        self.app = app
        self.template_path = getattr(app, 'selected_template_path', '')
        self.output_path = output_path
        self.update_doc_path = update_doc_path
        self.temp_dir = "temp_cropped_images"
        ensure_directory(self.temp_dir)
        
        self.performance = PerformanceSection(app, self.temp_dir)
        self.waveform = WaveformSection(app, self.temp_dir)

    def generate(self, progress_callback=None):
        doc_path = self.update_doc_path if self.update_doc_path and os.path.exists(self.update_doc_path) else self.template_path
        doc = Document(doc_path)
        
        perf_checked = [self.app.performancedata_list.item(i).text() for i in range(self.app.performancedata_list.count()) if self.app.performancedata_list.item(i).checkState() == Qt.Checked]
        wave_checked = [self.app.waveforms_list.item(i).text() for i in range(self.app.waveforms_list.count()) if self.app.waveforms_list.item(i).checkState() == Qt.Checked]

        if perf_checked:
            if progress_callback: progress_callback(70, "Writing Performance Data...")
            perf_data = self.performance.get_data(perf_checked)
            self.performance.add_section(doc, doc.element.body[-1], perf_checked, perf_data, None)

        if wave_checked:
            if progress_callback: progress_callback(85, "Writing Waveforms...")
            wave_files = self.waveform.get_images_with_custom_crop(wave_checked)
            self.waveform.add_section(doc, doc.element.body[-1], wave_checked, wave_files)

        if hasattr(self.app, 'pixl_file_path') and self.app.pixl_file_path:
            if progress_callback: progress_callback(90, "Appending PIXls Spreadsheet...")
            add_pixls_designer_table(doc, self.app.pixl_file_path)

        if hasattr(self.app, 'bom_file_path') and self.app.bom_file_path:
            if progress_callback: progress_callback(95, "Appending Bill of Materials...")
            add_bom_table(doc, self.app.bom_file_path)

        doc.save(self.output_path)
        os.startfile(self.output_path)
        remove_directory(self.temp_dir)