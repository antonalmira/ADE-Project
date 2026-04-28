import os
import re
from PyQt5.QtCore import Qt
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import log_message
from word_utils import add_caption_field, format_value_units
from image_utils import crop_and_save
from list_updater import waveform_testnames

class WaveformSection:
    def __init__(self, app, temp_dir):
        self.app = app
        self.temp_dir = temp_dir

    def get_first_two_words(self, filename):
        words = re.split(r'\s+|-|_', filename.lower())
        return ' '.join(words[:2]).strip()

    def get_images_with_custom_crop(self, waveform_items):
        """Gather and crop images based on UI settings."""
        try:
            c_left = int(self.app.left_input.text()) if self.app.left_input.text() else 0
            c_top = int(self.app.upper_input.text()) if self.app.upper_input.text() else 0
            c_right = int(self.app.right_input.text()) if self.app.right_input.text() else 0
            c_bottom = int(self.app.lower_input.text()) if self.app.lower_input.text() else 0
        except ValueError:
            c_left = c_top = c_right = c_bottom = 0

        waveform_folder = self.app.waveforms_path.text()
        if not waveform_folder or not os.path.isdir(waveform_folder):
            return {}

        checked_files_metadata = []
        current_test_category = None
        for i in range(self.app.available_data_list__waveforms.count()):
            item = self.app.available_data_list__waveforms.item(i)
            if not item: continue
            
            if item.text() in waveform_items:
                current_test_category = item.text()
            elif current_test_category and item.checkState() == Qt.Checked:
                checked_files_metadata.append((current_test_category, item.text(), item.data(Qt.UserRole)))

        waveform_files = {}
        for category in waveform_items:
            waveform_files[category] = []
            prefix = next((key for key, value in waveform_testnames.items() if value == category), '')
            
            for parent_cat, file_name, custom_cap in checked_files_metadata:
                if parent_cat == category:
                    if self.get_first_two_words(file_name) == prefix:
                        original_path = os.path.join(waveform_folder, file_name)
                        if os.path.exists(original_path):
                            cropped_path = crop_and_save(original_path, c_left, c_top, c_right, c_bottom, self.temp_dir)
                            if cropped_path:
                                waveform_files[category].append({
                                    'path': cropped_path, 
                                    'custom_cap': custom_cap
                                })
        return waveform_files

    def add_section(self, doc, last_element, waveform_items, waveform_files):
        """Writes the gathered images and captions into the Word document."""
        for item in waveform_items:
            log_message(f"Adding waveform subheader: {item}")
            new_para = doc.add_paragraph(item, style='Heading 2')
            new_para.runs[0].font.size = Pt(14)
            last_element.getparent().insert(last_element.getparent().index(last_element) + 1, new_para._element)
            last_element = new_para._element
            
            if item in waveform_files and waveform_files[item]:
                # Create a 2-column table for images
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(3.5)
                table.columns[1].width = Inches(3.5)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                current_row = 0
                current_col = 0
                
                for img_data in waveform_files[item]:
                    if current_col >= 2:
                        current_col = 0
                        current_row += 1
                        table.add_row()
                    
                    cell = table.cell(current_row, current_col)
                    cell_para = cell.paragraphs[0]
                    run = cell_para.add_run()
                    run.add_picture(img_data['path'], width=Inches(3.4))
                    
                    # --- Caption Logic ---
                    custom_cap = img_data.get('custom_cap')
                    main_cap_text = ""
                    if isinstance(custom_cap, dict):
                        main_cap_text = custom_cap.get('caption', '')
                    elif isinstance(custom_cap, str):
                        main_cap_text = custom_cap
                    
                    if not main_cap_text:
                        base_name = os.path.splitext(os.path.basename(img_data['path']))[0]
                        main_cap_text = format_value_units(base_name)

                    caption_cell = cell.add_paragraph()
                    add_caption_field(caption_cell, main_cap_text, "Figure")
                    caption_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # --- Extra Variable Lines (CH, Zoom, Meas) ---
                    if isinstance(custom_cap, dict):
                        for key in ['ch_info', 'zoom_info', 'meas_info']:
                            text_val = custom_cap.get(key, "")
                            if text_val:
                                extra_p = cell.add_paragraph()
                                extra_p.paragraph_format.left_indent = Inches(0.8)
                                extra_p.paragraph_format.space_after = Pt(2)
                                extra_p.paragraph_format.space_before = Pt(0)
                                
                                # Style Vripple/Blue color if it's ch_info
                                is_ch = (key == 'ch_info')
                                parts = re.split(r'(VRIPPLE|Vripple|V_RIPPLE|V_ripple)', text_val)
                                for part in parts:
                                    if part.lower().replace('_', '') == 'vripple':
                                        r1 = extra_p.add_run("V")
                                        r2 = extra_p.add_run("RIPPLE")
                                        r2.font.subscript = True
                                        if is_ch:
                                            r1.font.color.rgb = RGBColor(0, 0, 255)
                                            r2.font.color.rgb = RGBColor(0, 0, 255)
                                    else:
                                        r = extra_p.add_run(part)
                                        if is_ch: r.font.color.rgb = RGBColor(0, 0, 255)

                    current_col += 1
                
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, table._element)
                last_element = table._element
                
        return last_element