import os
import openpyxl
import re
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtCore import Qt

from utils import log_message
from excel_utils import extract_excel_table
from word_utils import add_styled_table, add_caption_field, format_value_units
from image_utils import crop_and_save
from list_updater import performancedata_testnames

class PerformanceSection:
    def __init__(self, app, temp_dir):
        self.app = app
        self.temp_dir = temp_dir

    def get_prefix(self, filename):
        return filename.split('_')[0].strip().lower()

    def get_data(self, performance_items):
        performance_folder = self.app.performancedata_path.text()
        charts_base_dir = os.path.join(performance_folder, "Performance Data Charts")
        performance_data = {}
        
        checked_metadata = []
        current_test_category = None
        for i in range(self.app.available_data_list_performance.count()):
            item = self.app.available_data_list_performance.item(i)
            if not item: continue
            
            if item.text() in performance_items:
                current_test_category = item.text()
            elif current_test_category and item.checkState() == Qt.Checked:
                checked_metadata.append((current_test_category, item.text(), item.data(Qt.UserRole)))

        for item_name in performance_items:
            performance_data[item_name] = {'charts': [], 'tables': []}
            prefix = next((key for key, value in performancedata_testnames.items() if value == item_name), '')
            item_chart_folder = os.path.join(charts_base_dir, f"{item_name} Charts")
            
            for parent_cat, file_name, custom_cap in checked_metadata:
                if parent_cat != item_name:
                    continue

                if os.path.isdir(item_chart_folder):
                    subfolder = os.path.splitext(file_name)[0]
                    subfolder_path = os.path.join(item_chart_folder, subfolder)
                    if os.path.isdir(subfolder_path):
                        for file in sorted(os.listdir(subfolder_path)):
                            if file.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                                performance_data[item_name]['charts'].append({
                                    'path': os.path.join(subfolder_path, file), 
                                    'custom_cap': custom_cap
                                })

                if self.get_prefix(file_name) == prefix:
                    file_path = os.path.join(performance_folder, file_name)
                    try:
                        wb = openpyxl.load_workbook(file_path, data_only=True)
                        ws = wb.active
                        table_data, merged_cells = extract_excel_table(ws)
                        if table_data:
                            performance_data[item_name]['tables'].append({
                                'file_name': file_name,
                                'data': table_data,
                                'merged_cells': merged_cells,
                                'custom_cap': custom_cap
                            })
                        wb.close()
                    except Exception as e:
                        log_message(f"Error extracting table from {file_name}: {str(e)}")
                        
        return performance_data

    def add_section(self, doc, last_element, performance_items, performance_data, efficiency_table):
        try:
            c_left = int(self.app.left_input.text()) if self.app.left_input.text() else 0
            c_top = int(self.app.upper_input.text()) if self.app.upper_input.text() else 0
            c_right = int(self.app.right_input.text()) if self.app.right_input.text() else 0
            c_bottom = int(self.app.lower_input.text()) if self.app.lower_input.text() else 0
        except:
            c_left = c_top = c_right = c_bottom = 0

        for item in performance_items:
            log_message(f"Adding performance subheader: {item}")
            new_para = doc.add_paragraph(item, style='Heading 2')
            new_para.runs[0].font.size = Pt(12)
            last_element.getparent().insert(last_element.getparent().index(last_element) + 1, new_para._element)
            last_element = new_para._element

            for chart_data in performance_data.get(item, {}).get('charts', []):
                image_path = chart_data['path']
                custom_cap = chart_data.get('custom_cap')
                
                cropped_path = crop_and_save(image_path, c_left, c_top, c_right, c_bottom, self.temp_dir)
                
                if cropped_path:
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(cropped_path, width=Inches(6.0))
                    last_element.getparent().insert(last_element.getparent().index(last_element) + 1, img_para._element)
                    last_element = img_para._element
                    
                    caption_text = self._get_caption_text(item, os.path.basename(image_path), custom_cap)
                    caption_para = doc.add_paragraph()
                    add_caption_field(caption_para, caption_text, "Figure")
                    last_element.getparent().insert(last_element.getparent().index(last_element) + 1, caption_para._element)
                    last_element = caption_para._element
                    
                    last_element = self._add_custom_metadata_lines(doc, last_element, custom_cap)

            for table_info in performance_data.get(item, {}).get('tables', []):
                table = add_styled_table(
                    doc, len(table_info['data']), len(table_info['data'][0]), 
                    table_info['data'], table_info['merged_cells'], 
                    header_color='#0078AB', font_name='Calibri', font_size=9
                )
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, table._element)
                last_element = table._element
                
                caption_text = self._get_caption_text(item, table_info['file_name'], table_info.get('custom_cap'))
                caption_para = doc.add_paragraph()
                add_caption_field(caption_para, caption_text, "Table")
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, caption_para._element)
                last_element = caption_para._element
                
                last_element = self._add_custom_metadata_lines(doc, last_element, table_info.get('custom_cap'))

        return last_element

    def _get_caption_text(self, item_name, file_name, custom_cap):
        if isinstance(custom_cap, dict) and custom_cap.get('caption'):
            return custom_cap['caption']
        if isinstance(custom_cap, str) and custom_cap:
            return custom_cap
        
        clean_name = os.path.splitext(file_name)[0].replace('_', ' ')
        return format_value_units(f"{item_name} - {clean_name}")

    def _add_custom_metadata_lines(self, doc, last_element, custom_cap):
        if not isinstance(custom_cap, dict):
            return last_element
            
        for key in ['ch_info', 'zoom_info', 'meas_info']:
            text_val = custom_cap.get(key, "")
            if text_val:
                p = doc.add_paragraph(text_val)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, p._element)
                last_element = p._element
        return last_element