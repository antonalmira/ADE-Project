from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.text import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_styled_table(doc, rows, cols, data, merged_cells=None, header_color='#0078AB', font_name='Calibri', font_size=9):
    if merged_cells is None:
        merged_cells = []
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = True

    for i in range(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = 1  # Middle vertical alignment
            cell_text = str(data[i][j]) if i < len(data) and j < len(data[i]) else ""
            # Remove trailing newlines and collapse multiple consecutive newlines to one
            cell_text = re.sub(r'\n+', '\n', cell_text.rstrip('\n\r')).strip()
            if cell_text:  # Only add paragraph if text is non-empty
                cell._element.clear_content()  # Clear any existing content
                paragraph = cell.add_paragraph()
                paragraph.text = cell_text
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if i == 0:  # Header row
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White font
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), header_color)
                        cell._element.get_or_add_tcPr().append(shading_elm)

    for min_row, min_col, max_row, max_col in merged_cells:
        table.cell(min_row, min_col).merge(table.cell(max_row, max_col))

    return table

def Figure(paragraph, bold=False):
    run = paragraph.add_run()
    if bold:
        run.font.bold = True
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

def Table(paragraph, bold=False):
    run = paragraph.add_run()
    if bold:
        run.font.bold = True
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

def add_caption_field(paragraph, caption_text, caption_type):
    paragraph.style = 'Caption'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Make the caption label and number bold (e.g., "Table 2")
    label_run = paragraph.add_run(f"{caption_type} ")
    label_run.font.bold = True
    if caption_type == "Figure":
        Figure(paragraph, bold=True)
    elif caption_type == "Table":
        Table(paragraph, bold=True)
    # Use an en-dash (–) with spaces and non-bold run for the rest of the caption
    caption_run = paragraph.add_run(f" – {caption_text}")
    caption_run.font.bold = False

def format_value_units(text: str) -> str:
    """
    Insert a space between numeric values and immediately following unit letters.
    Examples:
        "12V 3A 1.25A" -> "12 V 3 A 1.25 A"
    Handles letters and common unit symbols (%, °, µ, Ω). Leaves already-correct strings unchanged.
    """
    if not isinstance(text, str) or not text:
        return text
    # Insert a single space before a run of unit letters/symbols that immediately follow a digit.
    # e.g. "12V" -> "12 V", "1.25A" -> "1.25 A"
    return re.sub(r'(?<=\d)(?P<unit>[A-Za-z%°µΩ]+)', r' \g<unit>', text)