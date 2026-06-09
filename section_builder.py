import os
import openpyxl
import re
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from PyQt5.QtCore import Qt

from word_utils import add_styled_table, add_caption_field, format_value_units
from image_utils import crop_and_save


class SectionBuilder:
    def __init__(self, app, temp_dir):
        self.app = app
        self.temp_dir = temp_dir

    def build_all(self, doc):
        for i in range(self.app.unified_tree.topLevelItemCount()):
            top_node = self.app.unified_tree.topLevelItem(i)
            if top_node.checkState(0) == Qt.Unchecked:
                continue

            chapter_name = top_node.text(0).replace(" [FOLDER CROPPED]", "").strip()
            anchor = self._find_or_create_chapter(doc, chapter_name)
            self._wipe_section(doc, anchor)

            self._process_node(
                top_node, doc, anchor, current_level=1,
                source_type=top_node.data(0, Qt.UserRole + 12)
            )

    def _find_or_create_chapter(self, doc, chapter_name):
        target = chapter_name.lower()
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('Heading 1'):
                clean_p_text = re.sub(
                    r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)', '', p.text.lower()
                ).strip()
                if target in clean_p_text:
                    return p._element

        new_para = doc.add_paragraph(chapter_name, style='Heading 1')
        return new_para._element

    def _wipe_section(self, doc, start_element):
        parent = start_element.getparent()
        found_start = False
        elements_to_remove = []
        for child in parent:
            if not found_start:
                if child == start_element:
                    found_start = True
                continue
            if child.tag.endswith('p'):
                p = Paragraph(child, doc)
                if p.style and p.style.name == 'Heading 1':
                    break
            elements_to_remove.append(child)
        for el in elements_to_remove:
            parent.remove(el)

    def _process_node(self, node, doc, last_element, current_level, source_type):
        is_root = (node.parent() is None)
        current_anchor = last_element

        if not is_root and node.data(0, Qt.UserRole + 2) == "folder":
            clean_name = re.sub(
                r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)', '',
                node.text(0).replace(" [FOLDER CROPPED]", "").strip()
            )
            level = min(current_level + 1, 9)

            new_para = doc.add_paragraph(clean_name, style=f'Heading {level}')
            if level == 2 and new_para.runs:
                new_para.runs[0].font.size = Pt(14)
            current_anchor.getparent().insert(
                current_anchor.getparent().index(current_anchor) + 1, new_para._element
            )
            current_anchor = new_para._element

            if node.data(0, Qt.UserRole + 11):
                setup_para = doc.add_paragraph("Test Set-up", style=f'Heading {min(level + 1, 9)}')
                current_anchor.getparent().insert(
                    current_anchor.getparent().index(current_anchor) + 1, setup_para._element
                )
                current_anchor = self._add_test_setup_table(doc, setup_para._element, clean_name)

                res_para = doc.add_paragraph("Test Results", style=f'Heading {min(level + 1, 9)}')
                current_anchor.getparent().insert(
                    current_anchor.getparent().index(current_anchor) + 1, res_para._element
                )
                current_anchor = res_para._element

            current_level += 1

        # Gather files at this level
        files = [
            node.child(i) for i in range(node.childCount())
            if node.child(i).checkState(0) != Qt.Unchecked
            and node.child(i).data(0, Qt.UserRole + 2) == "file"
        ]

        if files:
            # If Excel Source: Process Tables normally, Process Images 1-col wide
            if source_type == "excel":
                for child in files:
                    current_anchor = self._render_excel_item(child, doc, current_anchor)
            # If Folder Source: Render Images efficiently side-by-side (2-col)
            elif source_type == "folder":
                current_anchor = self._render_image_table(files, doc, current_anchor)

        # Recurse Subfolders
        for i in range(node.childCount()):
            child = node.child(i)
            if child.checkState(0) != Qt.Unchecked and child.data(0, Qt.UserRole + 2) == "folder":
                current_anchor = self._process_node(
                    child, doc, current_anchor, current_level, source_type
                )

        return current_anchor

    def _render_excel_item(self, node, doc, last_element):
        item_type = node.data(0, Qt.UserRole + 7)
        clean_name = re.sub(
            r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)', '',
            node.text(0).replace(" [IMAGE CROPPED]", "").strip()
        )

        if item_type == "table":
            excel_path = node.data(0, Qt.UserRole + 8)
            sheet_name = node.data(0, Qt.UserRole + 9)
            voltage = node.data(0, Qt.UserRole + 10)

            if excel_path and os.path.exists(excel_path):
                split_tables = self._parse_split_table(excel_path, sheet_name)
                tb_data = split_tables.get(voltage) or split_tables.get("")

                # Robust safety check to avoid list index out of range
                if tb_data and tb_data.get('data'):
                    table = add_styled_table(
                        doc, len(tb_data['data']), len(tb_data['data'][0]),
                        tb_data['data'], tb_data['merged_cells'],
                        num_header_rows=tb_data.get('header_rows', 1),
                        widths=tb_data.get('widths', None)
                    )
                    last_element.getparent().insert(
                        last_element.getparent().index(last_element) + 1, table._element
                    )
                    last_element = table._element

                    cap_para = doc.add_paragraph()
                    add_caption_field(cap_para, format_value_units(clean_name), "Table")
                    last_element.getparent().insert(
                        last_element.getparent().index(last_element) + 1, cap_para._element
                    )
                    last_element = cap_para._element

        elif item_type == "graph" or not item_type:
            image_path = node.data(0, Qt.UserRole + 1)
            crop = self._get_crop(node)
            cropped_path = crop_and_save(
                image_path, crop['left'], crop['top'], crop['right'], crop['bottom'], self.temp_dir
            )
            if cropped_path:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(cropped_path, width=Inches(6.0))
                last_element.getparent().insert(
                    last_element.getparent().index(last_element) + 1, img_para._element
                )
                last_element = img_para._element

                cap_para = doc.add_paragraph()
                add_caption_field(cap_para, format_value_units(clean_name), "Figure")
                last_element.getparent().insert(
                    last_element.getparent().index(last_element) + 1, cap_para._element
                )
                last_element = cap_para._element

        return last_element

    def _render_image_table(self, file_nodes, doc, last_element):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.5)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r, c = 0, 0
        for node in file_nodes:
            original_path = node.data(0, Qt.UserRole + 1)
            crop = self._get_crop(node)
            cropped_path = crop_and_save(
                original_path, crop['left'], crop['top'], crop['right'], crop['bottom'], self.temp_dir
            )
            if not cropped_path:
                continue

            if c >= 2:
                c, r = 0, r + 1
                table.add_row()

            cell = table.cell(r, c)
            cell.paragraphs[0].add_run().add_picture(cropped_path, width=Inches(3.4))

            clean_name = re.sub(
                r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)', '',
                node.text(0).replace(" [IMAGE CROPPED]", "").strip()
            )
            caption_cell = cell.add_paragraph()
            add_caption_field(caption_cell, format_value_units(clean_name), "Figure")
            caption_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
            c += 1

        last_element.getparent().insert(
            last_element.getparent().index(last_element) + 1, table._element
        )
        return table._element

    def _add_test_setup_table(self, doc, last_element, clean_name):
        data = [
            ["Parameter", "Value"],
            ["Input Voltage", "85 VAC, 115 VAC, 230 VAC"],
            ["Output Voltage", "5 V, 9 V, 12 V, 15 V, 20 V"],
            ["Output Load", "50%, 100%"],
            ["Soak Time per Line", "15 minutes"],
            ["Integration time", "1 minute"],
            ["Output Voltage Measurement", "On-board"]
        ]
        table = add_styled_table(
            doc, len(data), 2, data,
            header_color='#5DA7E9', font_name='Calibri', font_size=10, widths=[2.5, 3.5]
        )
        for p in table.rows[1].cells[0].paragraphs:
            if p.runs:
                p.runs[0].font.bold = True
        last_element.getparent().insert(
            last_element.getparent().index(last_element) + 1, table._element
        )
        cap_para = doc.add_paragraph()
        add_caption_field(cap_para, f"{clean_name} Test Set-up", "Table")
        table._element.addnext(cap_para._element)
        return cap_para._element

    def _get_crop(self, node):
        while node:
            if node.data(0, Qt.UserRole + 3):
                return node.data(0, Qt.UserRole + 3)
            node = node.parent()
        return {'left': 0, 'top': 0, 'right': 0, 'bottom': 0}

    # =========================================================================
    # UNREDACTED TABLE PARSING LOGIC
    # =========================================================================
    def _parse_split_table(self, file_path, sheet_name):
        wb = None
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            if sheet_name not in wb.sheetnames:
                return {}

            ws = wb[sheet_name]
            raw_data = [list(row) for row in ws.iter_rows(values_only=True)]
            
            # --- FIX: Ensure perfect rectangular grid to prevent IndexError on jagged rows ---
            if raw_data:
                max_cols = max((len(row) for row in raw_data), default=0)
                raw_data = [row + [None] * (max_cols - len(row)) for row in raw_data]
            # -------------------------------------------------------------------------------

            raw_merges = [
                (m.min_row - 1, m.min_col - 1, m.max_row - 1, m.max_col - 1)
                for m in ws.merged_cells.ranges
            ]
        finally:
            if wb is not None:
                try:
                    wb.close()
                except Exception:
                    pass

        valid_rows = [
            i for i, row in enumerate(raw_data)
            if any(c is not None and str(c).strip() != "" for c in row)
        ]
        if not valid_rows:
            return {}

        valid_cols = []
        for j in range(len(raw_data[0])):
            col_cells = [raw_data[i][j] for i in valid_rows]
            is_empty = not any(c is not None and str(c).strip() != "" for c in col_cells)
            is_margin_col = any(c is not None and "margin" in str(c).lower() for c in col_cells)
            if not is_empty and not is_margin_col:
                valid_cols.append(j)

        if not valid_cols:
            return {}

        row_mapping = {old_i: new_i for new_i, old_i in enumerate(valid_rows)}
        col_mapping = {old_j: new_j for new_j, old_j in enumerate(valid_cols)}

        trimmed_data = []
        for old_i in valid_rows:
            trimmed_data.append([raw_data[old_i][old_j] for old_j in valid_cols])

        trimmed_merges = []
        for r_min, c_min, r_max, c_max in raw_merges:
            new_r_min = next((row_mapping[r] for r in range(r_min, r_max + 1) if r in row_mapping), None)
            new_r_max = next((row_mapping[r] for r in range(r_max, r_min - 1, -1) if r in row_mapping), None)
            new_c_min = next((col_mapping[c] for c in range(c_min, c_max + 1) if c in col_mapping), None)
            new_c_max = next((col_mapping[c] for c in range(c_max, c_min - 1, -1) if c in col_mapping), None)

            if (new_r_min is not None and new_r_max is not None
                    and new_c_min is not None and new_c_max is not None):
                if new_r_min <= new_r_max and new_c_min <= new_c_max:
                    trimmed_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))

        raw_data = trimmed_data
        raw_merges = trimmed_merges

        group_col_idx = -1
        header_end_idx = -1

        header_keywords = ['vac', 'vin', 'v_in', 'input', 'input voltage', 'input (vac)', 'line voltage']
        for i, row in enumerate(raw_data[:10]):
            for j, cell in enumerate(row):
                val = str(cell).strip().lower()
                if val in header_keywords or 'input voltage' in val or 'input (vac)' in val:
                    group_col_idx = j
                    break
            if group_col_idx != -1:
                break

        if group_col_idx == -1:
            safe_pattern = re.compile(
                r'^(85|90|100|115|132|180|230|264|265|277)(\.0)?\s*(vac|v)?$', re.IGNORECASE
            )
            for i, row in enumerate(raw_data):
                for j, cell in enumerate(row):
                    val = str(cell).strip() if cell is not None else ""
                    if safe_pattern.match(val):
                        group_col_idx = j
                        header_end_idx = i
                        break
                if group_col_idx != -1:
                    break

        if group_col_idx == -1:
            return {"": {"data": self._format_data(raw_data), "merged_cells": raw_merges, "header_rows": 2}}

        generic_pattern = re.compile(r'^(\d{2,3}(\.\d+)?)\s*(vac|v)?$', re.IGNORECASE)
        if header_end_idx == -1:
            for i in range(len(raw_data)):
                val = str(raw_data[i][group_col_idx]).strip() if raw_data[i][group_col_idx] is not None else ""
                if generic_pattern.match(val):
                    header_end_idx = i
                    break

        if header_end_idx == -1:
            return {"": {"data": self._format_data(raw_data), "merged_cells": raw_merges, "header_rows": 2}}

        headers = raw_data[:header_end_idx]
        groups = {}
        
        # FIX: Ensure current_group defaults to an empty string instead of None so appending is never skipped
        current_group = ""
        for i in range(header_end_idx, len(raw_data)):
            row = raw_data[i]
            if not any(v is not None for v in row):
                continue

            val = str(row[group_col_idx]).strip() if row[group_col_idx] is not None else ""
            match = generic_pattern.match(val)
            if match:
                current_group = match.group(1)

            if current_group not in groups:
                groups[current_group] = []
                
            groups[current_group].append((i, row))

        split_tables = {}
        is_eff_table = "eff" in sheet_name.lower()

        for group_name, rows_with_idx in groups.items():
            # Robust guard: ensuring empty dictionary lists don't crash the index lookups
            if not rows_with_idx:
                continue
                
            if is_eff_table:
                group_data = [
                    ["Load", "Design Performance", "", "", "", "", "", "Efficiency Standards", "", "Remarks"],
                    ["", "PIN", "VOUT at\nPCB", "IOUT", "POUT", "Efficiency\nat PCB", "Average\nEfficiency", "DOE6\nLimit", "CoC v5\nTier 2", ""],
                    ["(A)", "(W)", "(VDC)", "(mADC)", "(W)", "(%)", "(%)", "(%)", "(%)", ""]
                ]
                group_merges = [(0, 0, 1, 0), (0, 1, 0, 6), (0, 7, 0, 8), (0, 9, 2, 9)]

                start_old_row_idx = rows_with_idx[0][0]
                for old_idx, row in rows_with_idx:
                    clean_row = list(row[:group_col_idx] + row[group_col_idx + 1:])
                    while len(clean_row) < 10:
                        clean_row.append("")
                    clean_row = clean_row[:10]

                    load_val = str(clean_row[0]).strip()
                    if load_val.isdigit() or load_val.replace('.', '', 1).isdigit():
                        val = float(load_val)
                        clean_row[0] = f"{int(val)}%" if val.is_integer() else f"{val}%"

                    clean_row[9] = ""
                    group_data.append(clean_row)

                for r_min, c_min, r_max, c_max in raw_merges:
                    new_c_min = c_min if c_min <= group_col_idx else c_min - 1
                    new_c_max = c_max if c_max < group_col_idx else c_max - 1

                    if c_min == group_col_idx and c_max == group_col_idx:
                        continue

                    if r_min >= start_old_row_idx and r_max <= rows_with_idx[-1][0]:
                        new_r_min = (r_min - start_old_row_idx) + 3
                        new_r_max = (r_max - start_old_row_idx) + 3

                        if new_c_min <= 9 <= new_c_max:
                            continue

                        if new_c_max >= new_c_min:
                            group_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))

                split_tables[group_name] = {
                    "data": self._format_data(group_data),
                    "merged_cells": group_merges,
                    "header_rows": 3,
                    "widths": [0.6, 0.55, 0.7, 0.6, 0.55, 0.7, 0.7, 0.6, 0.6, 0.8]
                }
            else:
                group_data = []
                group_merges = []

                for h_row in headers:
                    group_data.append(h_row[:group_col_idx] + h_row[group_col_idx + 1:])

                start_old_row_idx = rows_with_idx[0][0]
                for old_idx, row in rows_with_idx:
                    group_data.append(row[:group_col_idx] + row[group_col_idx + 1:])

                for r_min, c_min, r_max, c_max in raw_merges:
                    new_c_min = c_min if c_min <= group_col_idx else c_min - 1
                    new_c_max = c_max if c_max < group_col_idx else c_max - 1

                    if c_min == group_col_idx and c_max == group_col_idx:
                        continue

                    if r_max < header_end_idx:
                        if new_c_max >= new_c_min:
                            group_merges.append((r_min, new_c_min, r_max, new_c_max))
                    elif r_min >= start_old_row_idx and r_max <= rows_with_idx[-1][0]:
                        new_r_min = (r_min - start_old_row_idx) + len(headers)
                        new_r_max = (r_max - start_old_row_idx) + len(headers)
                        if new_c_max >= new_c_min:
                            group_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))

                split_tables[group_name] = {
                    "data": self._format_data(group_data),
                    "merged_cells": group_merges,
                    "header_rows": len(headers),
                    "widths": None
                }

        return split_tables

    def _format_data(self, data):
        formatted = []
        for row in data:
            new_row = []
            for cell in row:
                if isinstance(cell, float):
                    if cell == 0.0:
                        new_row.append("0")
                    else:
                        try:
                            # 6 significant figures logic
                            val_str = f"{cell:.6g}"
                            if 'e' in val_str.lower():
                                val_str = f"{float(val_str):g}"
                            new_row.append(val_str)
                        except Exception:
                            new_row.append(str(cell))
                elif cell is None:
                    new_row.append("")
                else:
                    new_row.append(str(cell))
            formatted.append(new_row)
        return formatted