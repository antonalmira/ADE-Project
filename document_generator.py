import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.shared import OxmlElement, qn
from PyQt5.QtCore import Qt

import win32com.client
import pythoncom

from performance_section import PerformanceSection
from waveform_section import WaveformSection
from utils import ensure_directory, remove_directory


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_inner_borders(table, hex_color):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_table_all_borders(table, hex_color):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def format_text_specs(text):
    if not isinstance(text, str):
        return text
    text = re.sub(r'(?<=\d)(?=[a-zA-ZµΩ°])', ' ', text)
    for term in ['vac', 'vdc', 'vor', 'kp']:
        text = re.sub(fr'\b{term}\b', term.upper(), text, flags=re.IGNORECASE)
    return text


def apply_column_widths(table, width_inches_list):
    for row in table.rows:
        for idx, width in enumerate(width_inches_list):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def find_heading(doc, text_to_find):
    target = text_to_find.lower()
    for p in doc.paragraphs:
        if target in p.text.lower() and 'toc' not in p.style.name.lower():
            if 'heading' in p.style.name.lower():
                return p
    for p in doc.paragraphs:
        if target in p.text.lower() and 'toc' not in p.style.name.lower():
            if len(p.text) < 100:
                return p
    return None


class DocGenerator:
    def __init__(self, app, output_path):
        self.app = app
        self.template_path = getattr(app, 'selected_template_path', '')
        self.output_path = output_path
        self.temp_dir = "temp_cropped_images"
        ensure_directory(self.temp_dir)

        self.performance = PerformanceSection(app, self.temp_dir)
        self.waveform = WaveformSection(app, self.temp_dir)

    def _update_word_fields(self, filepath):
        """Opens Word to force-update all SEQ Fields (Figure/Table numbers) safely."""
        # CoInitialize must be called on whatever thread invokes COM objects.
        # This method may be called from the main thread (after the worker finishes),
        # so we initialize and uninitialize COM explicitly here.
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")

            # Open Word visible but minimized so the user can interact with any
            # login prompts (e.g. SharePoint) without causing an infinite hang.
            word.Visible = True
            try:
                word.WindowState = 2  # wdWindowStateMinimize
            except Exception:
                pass

            word.DisplayAlerts = False
            word.AutomationSecurity = 3  # Disable macros and Protected View warnings

            doc_com = word.Documents.Open(
                os.path.abspath(filepath),
                ConfirmConversions=False,
                ReadOnly=False,
                AddToRecentFiles=False
            )

            word.Selection.WholeStory()
            word.Selection.Fields.Update()

            for toc in doc_com.TablesOfContents:
                toc.Update()

            for tof in doc_com.TablesOfFigures:
                tof.Update()

            # Decouple Save and Close to prevent "Save As" / SharePoint sync prompts
            doc_com.Save()
            doc_com.Close(SaveChanges=False)

        except Exception as e:
            print(f"Non-critical Error: Could not automatically update Word fields: {e}")
        finally:
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def generate(self, progress_callback=None):
        doc = Document(self.template_path)

        if self.app.performance_tree.topLevelItemCount() > 0:
            if progress_callback:
                progress_callback(70, "Writing Performance Data...")
            self.performance.add_section(doc, doc.element.body[-1])

        if self.app.waveform_tree.topLevelItemCount() > 0:
            if progress_callback:
                progress_callback(85, "Writing Waveforms...")
            self.waveform.add_section(doc, doc.element.body[-1])

        if getattr(self.app, 'bom_file_path', None) and os.path.exists(self.app.bom_file_path):
            if progress_callback:
                progress_callback(90, "Appending Bill of Materials...")
            add_bom_table(doc, self.app.bom_file_path)

        if getattr(self.app, 'pixls_file_path', None) and os.path.exists(self.app.pixls_file_path):
            if progress_callback:
                progress_callback(95, "Appending Design Spreadsheet...")
            add_pixls_designer_table(doc, self.app.pixls_file_path)

        doc.save(self.output_path)

        if progress_callback:
            progress_callback(98, "Calculating Figure and Table numbers...")
        self._update_word_fields(self.output_path)

        os.startfile(self.output_path)
        remove_directory(self.temp_dir)


def add_bom_table(document, excel_path):
    try:
        xl = pd.ExcelFile(excel_path)
        elec_sheet = next((s for s in xl.sheet_names if 'elec' in s.lower()), None)
        mech_sheet = next((s for s in xl.sheet_names if 'mech' in s.lower()), None)

        if not elec_sheet and not mech_sheet:
            elec_sheet = next(
                (s for s in xl.sheet_names if 'bom' in s.lower()), xl.sheet_names[0]
            )

        sheets_to_process = []
        if elec_sheet:
            sheets_to_process.append((elec_sheet, "Electrical BOM"))
        if mech_sheet:
            sheets_to_process.append((mech_sheet, "Mechanical BOM"))

        main_bom_anchor = find_heading(document, "Bill of Materials")
        last_fallback_node = main_bom_anchor._element if main_bom_anchor else None

        for sheet_name, target_heading in sheets_to_process:
            df_raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
            header_idx = 0
            for i, row in df_raw.iterrows():
                row_str = [str(x).upper() for x in row.values if not pd.isna(x)]
                if any(k in s for s in row_str for k in ['DESIGNATOR', 'DESCRIPTION', 'PART NUMBER']):
                    header_idx = i
                    break

            df = pd.read_excel(excel_path, sheet_name=sheet_name, skiprows=header_idx)
            bom_columns = [
                'Item', 'Quantity', 'Designator', 'Value', 'Description',
                'Manufacturer Part Number', 'Manufacturer'
            ]

            rename_map = {}
            for col in df.columns:
                col_name = str(col).lower()
                # Use word-boundary matching to avoid false positives on substrings.
                # e.g. 'val' alone previously matched 'Interval', 'Approval', etc.
                if 'part number' in col_name and 'man' in col_name:
                    rename_map[col] = 'Manufacturer Part Number'
                elif re.search(r'\bman(ufacturer)?\b', col_name) and 'part' not in col_name:
                    rename_map[col] = 'Manufacturer'
                elif re.search(r'\bqty\b|\bquantity\b', col_name):
                    rename_map[col] = 'Quantity'
                elif re.search(r'\bdesc(ription)?\b', col_name):
                    rename_map[col] = 'Description'
                elif re.search(r'\bdesig(nator)?\b|\bref\b', col_name):
                    rename_map[col] = 'Designator'
                elif re.search(r'\bval(ue)?\b', col_name):
                    rename_map[col] = 'Value'
                elif re.search(r'\bitem\b', col_name):
                    rename_map[col] = 'Item'

            df = df.rename(columns=rename_map)

            existing_cols = [col for col in bom_columns if col in df.columns]
            df = df[existing_cols].dropna(how='all')
            if df.empty:
                continue

            table = document.add_table(rows=1, cols=len(df.columns))
            set_table_inner_borders(table, 'C0C0C0')
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            table.rows[0].height = Pt(35)

            for i, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[i]
                cell.text = str(col_name)
                set_cell_background(cell, '0085CA')
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    cell_val = "" if pd.isna(value) else str(value)
                    col_name = df.columns[i]

                    if col_name == 'Designator':
                        cell_val = cell_val.replace(',', '\n').replace(' ', '')
                    else:
                        cell_val = format_text_specs(cell_val)

                    cell = row_cells[i]
                    cell.text = cell_val
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if col_name == 'Description'
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)

            col_widths = {
                'Item': 0.4, 'Quantity': 0.6, 'Designator': 0.72, 'Value': 0.6,
                'Description': 2.16, 'Manufacturer Part Number': 1.65, 'Manufacturer': 1.01
            }
            widths = []
            for c in df.columns:
                w = col_widths.get(c, 1.0)
                if c == 'Description' and 'Value' not in df.columns:
                    w += 0.6
                widths.append(w)

            apply_column_widths(table, widths)

            anchor = find_heading(document, target_heading)
            if anchor:
                anchor._element.addnext(table._element)
            else:
                if last_fallback_node is not None:
                    heading = document.add_paragraph(target_heading, style='Heading 3')
                    last_fallback_node.addnext(heading._element)
                    heading._element.addnext(table._element)
                    last_fallback_node = table._element
                else:
                    heading = document.add_paragraph(target_heading, style='Heading 2')
                    document.element.body.append(heading._element)
                    document.element.body.append(table._element)
                    last_fallback_node = table._element

    except Exception as e:
        print(f"BOM Error: {e}")


def add_pixls_designer_table(document, excel_path):
    try:
        xl = pd.ExcelFile(excel_path)
        sheet_name = next((s for s in xl.sheet_names if 'PIX' in s or 'Design' in s), None)
        if not sheet_name:
            return

        df_raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
        header_idx = 0
        for i, row in df_raw.iterrows():
            row_str = [str(x).upper() for x in row.values]
            if 'INPUT' in row_str and 'OUTPUT' in row_str:
                header_idx = i
                break

        df = pd.read_excel(excel_path, sheet_name=sheet_name, skiprows=header_idx).iloc[:, :7]
        clean_headers = ["" if "Unnamed" in str(col) else str(col) for col in df.columns.tolist()]
        df.columns = ['Row Num', 'Parameter Name', 'INPUT', 'INFO', 'OUTPUT', 'UNIT', 'Description']

        table = document.add_table(rows=1, cols=len(df.columns))
        set_table_all_borders(table, 'C0C0C0')

        for i, col_name in enumerate(clean_headers):
            cell = table.rows[0].cells[i]
            cell.text = format_text_specs(col_name)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, 'FFFFFF')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        for _, row in df.iterrows():
            if pd.isna(row['Parameter Name']) and pd.isna(row['INPUT']):
                continue
            row_cells = table.add_row().cells
            is_subheader = pd.isna(row['INPUT']) and pd.isna(row['OUTPUT']) and pd.isna(row['UNIT'])

            if is_subheader:
                cell_0 = row_cells[0]
                num_val = str(row['Row Num']) if not pd.isna(row['Row Num']) else ""
                if num_val.endswith('.0'):
                    num_val = num_val[:-2]
                cell_0.text = num_val
                cell_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p0 = cell_0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r0 = p0.runs[0]
                r0.font.name = 'Calibri'
                r0.font.size = Pt(8)

                main_cell = row_cells[1]
                main_cell.merge(row_cells[5])
                main_cell.text = format_text_specs(str(row['Parameter Name']))
                set_cell_background(main_cell, 'D9D9D9')
                main_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = main_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

                desc_cell = row_cells[6]
                desc_val = "" if pd.isna(row['Description']) else format_text_specs(str(row['Description']))
                desc_cell.text = desc_val
                set_cell_background(desc_cell, 'FFFFFF')
                desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_desc = desc_cell.paragraphs[0]
                p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p_desc.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 0, 0)

            else:
                for i, col_name in enumerate(df.columns):
                    val = "" if pd.isna(row[col_name]) else format_text_specs(str(row[col_name]))
                    if col_name == 'Row Num' and val.endswith('.0'):
                        val = val[:-2]

                    cell = row_cells[i]
                    cell.text = val
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if col_name in ['Parameter Name', 'Description']
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        apply_column_widths(table, [0.32, 1.37, 0.7, 0.7, 0.7, 0.5, 2.52])

        anchor = find_heading(document, "Design Spreadsheet")
        if anchor:
            anchor._element.addnext(table._element)
        else:
            heading = document.add_paragraph("Design Spreadsheet", style='Heading 2')
            heading._element.addnext(table._element)

    except Exception as e:
        print(f"PIXls Error: {e}")