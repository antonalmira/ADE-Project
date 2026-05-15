import os
import openpyxl
import re
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from PyQt5.QtCore import Qt

from utils import log_message, ensure_directory
from word_utils import add_styled_table, add_caption_field, format_value_units
from image_utils import crop_and_save

class PerformanceSection:
    def __init__(self, app, temp_dir):
        self.app = app
        self.temp_dir = temp_dir
        self.perf_start_idx = 0

    def _strip_chapter_numbering(self, text):
        pattern = r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)(?![vV][aA]?[cCdD]?\b|[aA]\b|[mM][aA]\b|[wW]\b|[hH][zZ]\b|%)'
        return re.sub(pattern, '', text).strip()

    def add_section(self, doc, last_element, _=None):
        self.perf_start_idx, main_perf_anchor = self._get_main_performance_anchor(doc)
        
        if main_perf_anchor: 
            last_element = main_perf_anchor._element
            self._wipe_template_performance_section(doc, last_element)
        else:
            new_para = doc.add_paragraph("Performance Data", style='Heading 1')
            last_element.getparent().insert(last_element.getparent().index(last_element) + 1, new_para._element)
            last_element = new_para._element
            self.perf_start_idx = len(doc.paragraphs) - 1

        root = self.app.performance_tree.invisibleRootItem()
        return self._process_node(root, doc, last_element)

    def _get_main_performance_anchor(self, doc):
        for idx, p in enumerate(doc.paragraphs):
            if p.style and p.style.name.startswith('Heading'):
                clean_p_text = self._strip_chapter_numbering(p.text).lower()
                if "performance data" in clean_p_text:
                    return idx, p
        return 0, None

    def _wipe_template_performance_section(self, doc, start_element):
        parent = start_element.getparent()
        found_start = False
        elements_to_remove = []
        for child in parent:
            if not found_start:
                if child == start_element: found_start = True
                continue
            if child.tag.endswith('p'):
                p = Paragraph(child, doc)
                if p.style and p.style.name == 'Heading 1': break
            elements_to_remove.append(child)
        for el in elements_to_remove: parent.remove(el)

    def _process_node(self, node, doc, last_element):
        is_root = (node == self.app.performance_tree.invisibleRootItem())

        if not is_root and node.checkState(0) == Qt.Unchecked: return last_element

        current_anchor = last_element

        if not is_root and node.data(0, Qt.UserRole + 2) == "folder":
            clean_name = self._strip_chapter_numbering(node.text(0).replace(" [CROP SET]", ""))
            
            # Print Category Heading
            new_para = doc.add_paragraph(clean_name, style='Heading 2')
            current_anchor.getparent().insert(current_anchor.getparent().index(current_anchor) + 1, new_para._element)
            current_anchor = new_para._element 
            
            # Check if user enabled the Setup Table for this folder
            include_setup = node.data(0, Qt.UserRole + 11)
            if include_setup is None: include_setup = False
            
            if include_setup:
                setup_para = doc.add_paragraph("Test Set-up", style='Heading 3')
                current_anchor.getparent().insert(current_anchor.getparent().index(current_anchor) + 1, setup_para._element)
                current_anchor = setup_para._element
                
                # Automatically build the standard test setup table
                current_anchor = self._add_test_setup_table(doc, current_anchor, clean_name)
            
            res_para = doc.add_paragraph("Test Results", style='Heading 3')
            current_anchor.getparent().insert(current_anchor.getparent().index(current_anchor) + 1, res_para._element)
            current_anchor = res_para._element

        # Iterate children
        for i in range(node.childCount()):
            child = node.child(i)
            if child.checkState(0) == Qt.Unchecked: continue
            
            if child.data(0, Qt.UserRole + 2) == "file":
                current_anchor = self._render_file_item(child, doc, current_anchor)
            elif child.data(0, Qt.UserRole + 2) == "folder":
                current_anchor = self._process_node(child, doc, current_anchor)

        return current_anchor

    def _add_test_setup_table(self, doc, last_element, clean_name):
        data = [
            ["Parameter", "Value"],
            ["Input Voltage", "85 VAC, 115 VAC, 230 VAC"],
            ["Output Voltage", "5 V, 9 V, 12 V, 15 V, 20 V"],
            ["Output Load", "50%, 100%"],
            ["Soak Time per Line", "15 minutes"],
            ["Integration time", "1 minute"],
            ["Output Voltage Measurement", "On-board"]
        ]
        
        table = add_styled_table(
            doc, len(data), 2, data, 
            header_color='#5DA7E9', font_name='Calibri', font_size=10,
            num_header_rows=1, widths=[2.5, 3.5]
        )
        
        for r_idx, row in enumerate(table.rows):
            if r_idx == 0: continue
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    
        last_element.getparent().insert(last_element.getparent().index(last_element) + 1, table._element)
        last_element = table._element
        
        cap_para = doc.add_paragraph()
        add_caption_field(cap_para, f"{clean_name} Test Set-up", "Table")
        last_element.getparent().insert(last_element.getparent().index(last_element) + 1, cap_para._element)
        
        return cap_para._element

    def _render_file_item(self, node, doc, last_element):
        item_type = node.data(0, Qt.UserRole + 7) 
        excel_path = node.data(0, Qt.UserRole + 8)
        sheet_name = node.data(0, Qt.UserRole + 9)
        voltage = node.data(0, Qt.UserRole + 10)
        custom_cap = node.data(0, Qt.UserRole)
        
        clean_name = self._strip_chapter_numbering(node.text(0).replace(" [CROP SET]", ""))

        if item_type == "table":
            if excel_path and os.path.exists(excel_path):
                split_tables = self._parse_split_table(excel_path, sheet_name)
                
                tb_data = None
                if voltage and voltage in split_tables: tb_data = split_tables[voltage]
                elif "" in split_tables: tb_data = split_tables[""]
                
                if tb_data:
                    table = add_styled_table(
                        doc, len(tb_data['data']), len(tb_data['data'][0]), 
                        tb_data['data'], tb_data['merged_cells'], 
                        header_color='#0085CA', font_name='Calibri', font_size=9,
                        num_header_rows=tb_data.get('header_rows', 1),
                        widths=tb_data.get('widths', None) 
                    )
                    last_element.getparent().insert(last_element.getparent().index(last_element) + 1, table._element)
                    last_element = table._element
                    
                    cap_para = doc.add_paragraph()
                    add_caption_field(cap_para, format_value_units(clean_name), "Table")
                    last_element.getparent().insert(last_element.getparent().index(last_element) + 1, cap_para._element)
                    last_element = cap_para._element

        elif item_type == "graph":
            image_path = node.data(0, Qt.UserRole + 1)
            crop_data = self._get_crop_for_node(node)
            c_l = int(crop_data.get('left', 0)); c_t = int(crop_data.get('top', 0))
            c_r = int(crop_data.get('right', 0)); c_b = int(crop_data.get('bottom', 0))

            cropped_path = crop_and_save(image_path, c_l, c_t, c_r, c_b, self.temp_dir)
            if cropped_path:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(cropped_path, width=Inches(6.0))
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, img_para._element)
                last_element = img_para._element
                
                cap_para = doc.add_paragraph()
                add_caption_field(cap_para, format_value_units(clean_name), "Figure")
                last_element.getparent().insert(last_element.getparent().index(last_element) + 1, cap_para._element)
                last_element = cap_para._element

        if isinstance(custom_cap, dict):
            for key in ['ch_info', 'zoom_info', 'meas_info']:
                text_val = custom_cap.get(key, "")
                if text_val:
                    p = doc.add_paragraph(text_val)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(0)
                    last_element.getparent().insert(last_element.getparent().index(last_element) + 1, p._element)
                    last_element = p._element

        return last_element

    def _get_crop_for_node(self, node):
        current = node
        while current:
            data = current.data(0, Qt.UserRole + 3)
            if data: return data
            current = current.parent()
        return {'left': '0', 'top': '0', 'right': '0', 'bottom': '0'}

    def _parse_split_table(self, file_path, sheet_name):
        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb[sheet_name]
        
        raw_data = [list(row) for row in ws.iter_rows(values_only=True)]
        raw_merges = [(m.min_row - 1, m.min_col - 1, m.max_row - 1, m.max_col - 1) for m in ws.merged_cells.ranges]
        wb.close()

        valid_rows = [i for i, row in enumerate(raw_data) if any(c is not None and str(c).strip() != "" for c in row)]
        if not valid_rows: return {}

        valid_cols = []
        for j in range(len(raw_data[0])):
            col_cells = [raw_data[i][j] for i in valid_rows]
            is_empty = not any(c is not None and str(c).strip() != "" for c in col_cells)
            is_margin_col = any(c is not None and "margin" in str(c).lower() for c in col_cells)
            
            if not is_empty and not is_margin_col:
                valid_cols.append(j)

        if not valid_cols: return {}

        row_mapping = {old_i: new_i for new_i, old_i in enumerate(valid_rows)}
        col_mapping = {old_j: new_j for new_j, old_j in enumerate(valid_cols)}
        
        trimmed_data = []
        for old_i in valid_rows:
            trimmed_data.append([raw_data[old_i][old_j] for old_j in valid_cols])
            
        trimmed_merges = []
        for r_min, c_min, r_max, c_max in raw_merges:
            new_r_min = next((row_mapping[r] for r in range(r_min, r_max+1) if r in row_mapping), None)
            new_r_max = next((row_mapping[r] for r in range(r_max, r_min-1, -1) if r in row_mapping), None)
            new_c_min = next((col_mapping[c] for c in range(c_min, c_max+1) if c in col_mapping), None)
            new_c_max = next((col_mapping[c] for c in range(c_max, c_min-1, -1) if c in col_mapping), None)
            
            if new_r_min is not None and new_r_max is not None and new_c_min is not None and new_c_max is not None:
                if new_r_min <= new_r_max and new_c_min <= new_c_max:
                    trimmed_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))

        raw_data = trimmed_data
        raw_merges = trimmed_merges

        group_col_idx = -1
        header_end_idx = -1
        
        header_keywords = ['vac', 'vin', 'v_in', 'input', 'input voltage', 'input (vac)', 'line voltage']
        for i, row in enumerate(raw_data[:10]):
            for j, cell in enumerate(row):
                val = str(cell).strip().lower()
                if val in header_keywords or 'input voltage' in val or 'input (vac)' in val:
                    group_col_idx = j
                    break
            if group_col_idx != -1: break

        if group_col_idx == -1:
            safe_pattern = re.compile(r'^(85|90|100|115|132|180|230|264|265|277)(\.0)?\s*(vac|v)?$', re.IGNORECASE)
            for i, row in enumerate(raw_data):
                for j, cell in enumerate(row):
                    val = str(cell).strip() if cell is not None else ""
                    if safe_pattern.match(val):
                        group_col_idx = j
                        header_end_idx = i
                        break
                if group_col_idx != -1: break

        if group_col_idx == -1:
            return {"": { "data": self._format_data(raw_data), "merged_cells": raw_merges, "header_rows": 2 }}

        generic_pattern = re.compile(r'^(\d{2,3}(\.\d+)?)\s*(vac|v)?$', re.IGNORECASE)
        if header_end_idx == -1:
            for i in range(len(raw_data)):
                val = str(raw_data[i][group_col_idx]).strip() if raw_data[i][group_col_idx] is not None else ""
                if generic_pattern.match(val):
                    header_end_idx = i
                    break

        if header_end_idx == -1:
            return {"": { "data": self._format_data(raw_data), "merged_cells": raw_merges, "header_rows": 2 }}

        headers = raw_data[:header_end_idx]
        groups = {}
        current_group = None
        for i in range(header_end_idx, len(raw_data)):
            row = raw_data[i]
            if not any(v is not None for v in row): continue 
            
            val = str(row[group_col_idx]).strip() if row[group_col_idx] is not None else ""
            match = generic_pattern.match(val)
            if match: current_group = match.group(1)
            
            if current_group not in groups: groups[current_group] = []
            if current_group: groups[current_group].append((i, row))
                
        split_tables = {}
        is_eff_table = "eff" in sheet_name.lower()
        
        for group_name, rows_with_idx in groups.items():
            if is_eff_table:
                group_data = [
                    ["Load", "Design Performance", "", "", "", "", "", "Efficiency Standards", "", "Remarks"],
                    ["", "PIN", "VOUT at\nPCB", "IOUT", "POUT", "Efficiency\nat PCB", "Average\nEfficiency", "DOE6\nLimit", "CoC v5\nTier 2", ""],
                    ["(A)", "(W)", "(VDC)", "(mADC)", "(W)", "(%)", "(%)", "(%)", "(%)", ""]
                ]
                group_merges = [(0,0,1,0), (0,1,0,6), (0,7,0,8), (0,9,2,9)]
                
                start_old_row_idx = rows_with_idx[0][0]
                for old_idx, row in rows_with_idx:
                    clean_row = list(row[:group_col_idx] + row[group_col_idx+1:])
                    while len(clean_row) < 10: clean_row.append("")
                    clean_row = clean_row[:10]
                    
                    load_val = str(clean_row[0]).strip()
                    if load_val.isdigit() or load_val.replace('.', '', 1).isdigit():
                        val = float(load_val)
                        clean_row[0] = f"{int(val)}%" if val.is_integer() else f"{val}%"
                        
                    clean_row[9] = ""
                    group_data.append(clean_row)

                for r_min, c_min, r_max, c_max in raw_merges:
                    new_c_min = c_min if c_min <= group_col_idx else c_min - 1
                    new_c_max = c_max if c_max < group_col_idx else c_max - 1
                    
                    if c_min == group_col_idx and c_max == group_col_idx: continue
                        
                    if r_min >= start_old_row_idx and r_max <= rows_with_idx[-1][0]:
                        new_r_min = (r_min - start_old_row_idx) + 3
                        new_r_max = (r_max - start_old_row_idx) + 3
                        
                        if new_c_min <= 9 <= new_c_max: continue
                            
                        if new_c_max >= new_c_min:
                            group_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))
                    
                split_tables[group_name] = {
                    "data": self._format_data(group_data),
                    "merged_cells": group_merges,
                    "header_rows": 3,
                    "widths": [0.6, 0.55, 0.7, 0.6, 0.55, 0.7, 0.7, 0.6, 0.6, 0.8] 
                }
            else:
                group_data = []
                group_merges = []
                
                for h_row in headers:
                    group_data.append(h_row[:group_col_idx] + h_row[group_col_idx+1:])
                    
                start_old_row_idx = rows_with_idx[0][0]
                for old_idx, row in rows_with_idx:
                    group_data.append(row[:group_col_idx] + row[group_col_idx+1:])
                    
                for r_min, c_min, r_max, c_max in raw_merges:
                    new_c_min = c_min if c_min <= group_col_idx else c_min - 1
                    new_c_max = c_max if c_max < group_col_idx else c_max - 1
                    
                    if c_min == group_col_idx and c_max == group_col_idx: continue
                        
                    if r_max < header_end_idx:
                        if new_c_max >= new_c_min:
                            group_merges.append((r_min, new_c_min, r_max, new_c_max))
                    elif r_min >= start_old_row_idx and r_max <= rows_with_idx[-1][0]:
                        new_r_min = (r_min - start_old_row_idx) + len(headers)
                        new_r_max = (r_max - start_old_row_idx) + len(headers)
                        if new_c_max >= new_c_min:
                            group_merges.append((new_r_min, new_c_min, new_r_max, new_c_max))
                            
                split_tables[group_name] = {
                    "data": self._format_data(group_data), 
                    "merged_cells": group_merges,
                    "header_rows": len(headers),
                    "widths": None
                }
            
        return split_tables

    def _format_data(self, data):
        formatted = []
        for row in data:
            new_row = []
            for cell in row:
                if isinstance(cell, float):
                    if cell == 0.0: new_row.append("0")
                    else:
                        try:
                            val_str = f"{cell:.3g}"
                            if 'e' in val_str.lower(): val_str = f"{float(val_str):g}"
                            new_row.append(val_str)
                        except:
                            new_row.append(str(cell))
                elif cell is None: new_row.append("")
                else: new_row.append(str(cell))
            formatted.append(new_row)
        return formatted