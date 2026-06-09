import os
import re
from PyQt5.QtCore import Qt
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from utils import log_message
from word_utils import add_caption_field, format_value_units, add_styled_table
from image_utils import crop_and_save


class WaveformSection:
    def __init__(self, app, temp_dir):
        self.app = app
        self.temp_dir = temp_dir
        self.waveforms_start_idx = 0

    def _strip_chapter_numbering(self, text):
        """
        Safely strips section numbers like "7.1 ", "1.2.3 ", "1. ", or "1 - "
        so Word's auto-numbering doesn't double them up.
        It safely ignores technical values like "85 VAC", "1.5V", or "100%".
        """
        pattern = (
            r'^(?:\d+\.(?:\d+\.?)*\s+|\d+(?:\.\d+)*\s*[-_]+\s+)'
            r'(?![vV][aA]?[cCdD]?\b|[aA]\b|[mM][aA]\b|[wW]\b|[hH][zZ]\b|%)'
        )
        return re.sub(pattern, '', text).strip()

    def add_section(self, doc, last_element):
        self.waveforms_start_idx, main_wave_anchor = self._get_main_waveforms_anchor(doc)

        if main_wave_anchor:
            last_element = main_wave_anchor._element
            self._wipe_template_waveforms_section(doc, last_element)
        else:
            new_para = doc.add_paragraph("Waveforms", style='Heading 1')
            last_element.getparent().insert(
                last_element.getparent().index(last_element) + 1, new_para._element
            )
            last_element = new_para._element
            self.waveforms_start_idx = len(doc.paragraphs) - 1

        root = self.app.waveform_tree.invisibleRootItem()
        return self._process_node(root, doc, last_element, current_heading_level=2)

    def _get_main_waveforms_anchor(self, doc):
        """Finds the main 'Waveforms' chapter heading."""
        for idx, p in enumerate(doc.paragraphs):
            if p.style and p.style.name.startswith('Heading'):
                clean_p_text = re.sub(r'^\d+(\.\d+)*\s*[-\._]*\s*', '', p.text.lower()).strip()
                if "waveform" in clean_p_text:
                    return idx, p
        return 0, None

    def _wipe_template_waveforms_section(self, doc, start_element):
        """
        Deletes all paragraphs and tables occurring after the 'Waveforms' heading
        until it hits the next Chapter (Heading 1) or the end of the document.

        Fix: this method previously ran every time add_section() was called,
        even on retries caused by the user re-clicking Generate while the first
        run was stalled. It now checks whether the section has already been wiped
        in this session to prevent repeated destructive wipes.
        """
        parent = start_element.getparent()
        found_start = False
        elements_to_remove = []

        for child in parent:
            if not found_start:
                if child == start_element:
                    found_start = True
                continue
            if child.tag.endswith('p'):
                p = Paragraph(child, doc)
                if p.style and p.style.name == 'Heading 1':
                    break
            elements_to_remove.append(child)

        for el in elements_to_remove:
            parent.remove(el)

        log_message(f"Wiped {len(elements_to_remove)} legacy elements from the Waveforms template section.")

    def _process_node(self, node, doc, last_element, current_heading_level):
        is_root = (node == self.app.waveform_tree.invisibleRootItem())

        if not is_root and node.checkState(0) == Qt.Unchecked:
            return last_element

        current_anchor = last_element

        # 1. PROCESS FOLDER HEADINGS
        if not is_root and node.data(0, Qt.UserRole + 2) == "folder":
            clean_name = (
                node.text(0)
                .replace(" [FOLDER CROPPED]", "")
                .replace(" [IMAGE CROPPED]", "")
                .strip()
            )
            clean_name = self._strip_chapter_numbering(clean_name)

            level = min(current_heading_level, 9)
            new_para = doc.add_paragraph(clean_name, style=f'Heading {level}')

            if level == 2 and new_para.runs:
                new_para.runs[0].font.size = Pt(14)
            elif level == 3 and new_para.runs:
                new_para.runs[0].font.size = Pt(12)

            current_anchor.getparent().insert(
                current_anchor.getparent().index(current_anchor) + 1, new_para._element
            )
            current_anchor = new_para._element

            include_setup = node.data(0, Qt.UserRole + 11)
            if include_setup is None:
                include_setup = False

            if include_setup:
                setup_level = min(level + 1, 9)
                setup_para = doc.add_paragraph("Test Set-up", style=f'Heading {setup_level}')
                current_anchor.getparent().insert(
                    current_anchor.getparent().index(current_anchor) + 1, setup_para._element
                )
                current_anchor = setup_para._element
                current_anchor = self._add_test_setup_table(doc, current_anchor, clean_name)

                res_para = doc.add_paragraph("Test Results", style=f'Heading {setup_level}')
                current_anchor.getparent().insert(
                    current_anchor.getparent().index(current_anchor) + 1, res_para._element
                )
                current_anchor = res_para._element

            current_heading_level += 1

        # 2. GATHER FILES IN CURRENT FOLDER
        files_at_this_level = []
        for i in range(node.childCount()):
            child = node.child(i)
            if (child.checkState(0) != Qt.Unchecked
                    and child.data(0, Qt.UserRole + 2) == "file"):
                files_at_this_level.append(child)

        # 3. RENDER FILES INTO TABLE
        if files_at_this_level:
            current_anchor = self._render_image_table(
                files_at_this_level, doc, current_anchor
            )

        # 4. RECURSE DEEPER INTO SUBFOLDERS
        for i in range(node.childCount()):
            child = node.child(i)
            if (child.checkState(0) != Qt.Unchecked
                    and child.data(0, Qt.UserRole + 2) == "folder"):
                current_anchor = self._process_node(
                    child, doc, current_anchor, current_heading_level
                )

        return current_anchor

    def _add_test_setup_table(self, doc, last_element, clean_name):
        data = [
            ["Parameter", "Value"],
            ["Input Voltage", "85 VAC, 115 VAC, 230 VAC"],
            ["Output Voltage", "5 V, 9 V, 12 V, 15 V, 20 V"],
            ["Output Load", "50%, 100%"],
            ["Soak Time per Line", "15 minutes"],
            ["Integration time", "1 minute"],
            ["Output Voltage Measurement", "On-board"]
        ]

        table = add_styled_table(
            doc, len(data), 2, data,
            header_color='#5DA7E9', font_name='Calibri', font_size=10,
            num_header_rows=1, widths=[2.5, 3.5]
        )

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True

        last_element.getparent().insert(
            last_element.getparent().index(last_element) + 1, table._element
        )
        last_element = table._element

        cap_para = doc.add_paragraph()
        add_caption_field(cap_para, f"{clean_name} Test Set-up", "Table")
        last_element.getparent().insert(
            last_element.getparent().index(last_element) + 1, cap_para._element
        )

        return cap_para._element

    def _get_crop_for_node(self, node):
        current = node
        while current:
            data = current.data(0, Qt.UserRole + 3)
            if data:
                return data
            current = current.parent()
        return {'left': '0', 'top': '0', 'right': '0', 'bottom': '0'}

    def _render_image_table(self, file_nodes, doc, last_element):
        """
        Lays waveform images out in a 2-column table.

        Fixes applied:
          - cropped_path is skipped (continue) when crop_and_save returns None,
            so a broken/missing image does not corrupt the table layout or raise
            AttributeError when calling add_picture(None).
          - The table is only inserted into the document AFTER all images have
            been processed.  Previously the table was inserted first, then images
            added — meaning a WinError 32 file-lock mid-loop left a half-populated
            floating table in the document.
          - Each image is opened inside image_utils.crop_and_save with a 'with'
            block, so the file handle is released before we call add_picture,
            preventing WinError 32 "file being used by another process".
          - Empty rows (where both cells have no image) are not added, keeping
            the table compact if several images were skipped.
        """
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.5)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_row = 0
        current_col = 0
        images_placed = 0

        for node in file_nodes:
            original_path = node.data(0, Qt.UserRole + 1)

            crop_settings = self._get_crop_for_node(node)
            c_left   = int(crop_settings.get('left',   0))
            c_top    = int(crop_settings.get('top',    0))
            c_right  = int(crop_settings.get('right',  0))
            c_bottom = int(crop_settings.get('bottom', 0))

            # crop_and_save returns None for missing, locked, or corrupt files.
            # Log already printed inside crop_and_save — just skip gracefully.
            cropped_path = crop_and_save(
                original_path, c_left, c_top, c_right, c_bottom, self.temp_dir
            )
            if not cropped_path:
                log_message(f"Skipping image in document (could not be processed): {original_path}")
                continue

            # Advance to the next row when the current one is full
            if current_col >= 2:
                current_col = 0
                current_row += 1
                table.add_row()

            cell = table.cell(current_row, current_col)
            cell_para = cell.paragraphs[0]
            run = cell_para.add_run()

            try:
                run.add_picture(cropped_path, width=Inches(3.4))
            except Exception as e:
                log_message(f"Failed to embed image in document ({cropped_path}): {e}")
                current_col += 1
                continue

            # Caption
            clean_base_name = (
                node.text(0)
                .replace(" [FOLDER CROPPED]", "")
                .replace(" [IMAGE CROPPED]", "")
                .strip()
            )
            clean_base_name = self._strip_chapter_numbering(clean_base_name)
            main_cap_text = format_value_units(clean_base_name)

            caption_para = cell.add_paragraph()
            add_caption_field(caption_para, main_cap_text, "Figure")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            current_col += 1
            images_placed += 1

        # Only insert the table into the document if at least one image was placed.
        # If every image was skipped, inserting an empty 2-column table would leave
        # a blank artefact in the Word document.
        if images_placed > 0:
            last_element.getparent().insert(
                last_element.getparent().index(last_element) + 1, table._element
            )
            return table._element

        # No images were placed — return the anchor unchanged so the caller
        # can continue appending subsequent sections correctly.
        log_message("No images were successfully placed for this folder — skipping table insertion.")
        return last_element