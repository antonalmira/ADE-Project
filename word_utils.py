from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def set_table_borders(table):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '595959')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def apply_formatted_text(paragraph, text, is_header, font_name, font_size):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0

    if not text:
        return

    font_color = RGBColor(255, 255, 255) if is_header else None
    pattern = re.compile(r'(PIN|VOUT|IOUT|POUT|mADC)')
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue

        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if font_color:
            run.font.color.rgb = font_color
        if is_header:
            run.font.bold = True

        if part == 'PIN':
            run.text = 'P'
            sub = paragraph.add_run('IN')
        elif part == 'VOUT':
            run.text = 'V'
            sub = paragraph.add_run('OUT')
        elif part == 'IOUT':
            run.text = 'I'
            sub = paragraph.add_run('OUT')
        elif part == 'POUT':
            run.text = 'P'
            sub = paragraph.add_run('OUT')
        elif part == 'mADC':
            run.text = 'mA'
            sub = paragraph.add_run('DC')
        else:
            run.text = part
            continue

        sub.font.name = font_name
        sub.font.size = Pt(font_size)
        if font_color:
            sub.font.color.rgb = font_color
        if is_header:
            sub.font.bold = True
        sub.font.subscript = True


def add_styled_table(doc, rows, cols, data, merged_cells=None, header_color='#0085CA',
                     font_name='Calibri', font_size=9, num_header_rows=1, widths=None):
    if merged_cells is None:
        merged_cells = []

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_borders(table)

    # Perform merges FIRST to prevent text duplication
    for min_row, min_col, max_row, max_col in merged_cells:
        try:
            table.cell(min_row, min_col).merge(table.cell(max_row, max_col))
        except (IndexError, ValueError):
            # Only catch coordinate-out-of-range errors, not programming bugs
            pass

    # Track which merged cells have already been populated
    processed_cells = set()

    for i in range(rows):
        is_header = i < num_header_rows
        for j in range(cols):
            cell = table.cell(i, j)

            # Check the underlying XML element to accurately detect merged cells
            if cell._tc in processed_cells:
                continue

            processed_cells.add(cell._tc)
            cell.vertical_alignment = 1  # WD_ALIGN_VERTICAL.CENTER

            cell_text = str(data[i][j]) if i < len(data) and j < len(data[i]) else ""
            cell_text = re.sub(r'\n+', '\n', cell_text.rstrip('\n\r')).strip()

            # Clear native ghost paragraphs inside merged cells
            for p in cell.paragraphs:
                p.clear()

            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            if cell_text:
                apply_formatted_text(paragraph, cell_text, is_header, font_name, font_size)
            else:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

            if is_header:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                shading_elm.set(qn('w:fill'), header_color.lstrip('#'))
                cell._element.get_or_add_tcPr().append(shading_elm)

    if widths:
        set_column_widths(table, widths)

    return table


def Figure(paragraph, bold=False):
    run = paragraph.add_run()
    if bold:
        run.font.bold = True
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def Table(paragraph, bold=False):
    run = paragraph.add_run()
    if bold:
        run.font.bold = True
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def add_caption_field(paragraph, caption_text, caption_type):
    paragraph.style = 'Caption'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{caption_type} ")
    label_run.font.bold = True
    if caption_type == "Figure":
        Figure(paragraph, bold=True)
    elif caption_type == "Table":
        Table(paragraph, bold=True)
    caption_run = paragraph.add_run(f" – {caption_text}")
    caption_run.font.bold = False


def format_value_units(text: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    return re.sub(r'(?<=\d)(?P<unit>[A-Za-z%°µΩ]+)', r' \g<unit>', text)